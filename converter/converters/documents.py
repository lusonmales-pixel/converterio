"""
Конвертация документов: PDF, DOCX, TXT, HTML.
"""

import io
from pathlib import Path

from .base import ConversionError

# PDF -> DOCX
try:
    from pdf2docx import Converter as Pdf2DocxConverter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

# DOCX, TXT
try:
    from docx import Document
    from docx.shared import Pt, Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# PDF generation
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# HTML -> PDF
try:
    from xhtml2pdf import pisa
    XHTML2PDF_AVAILABLE = True
except ImportError:
    XHTML2PDF_AVAILABLE = False


def convert_document(source_path: str, source_fmt: str, target_fmt: str, output_dir: Path) -> str:
    """
    Конвертирует документ.
    """
    base = Path(source_path).stem
    out_path = output_dir / f'{base}.{target_fmt}'

    if source_fmt == 'pdf' and target_fmt == 'docx':
        return _pdf_to_docx(source_path, str(out_path))
    if source_fmt == 'docx' and target_fmt == 'pdf':
        return _docx_to_pdf(source_path, str(out_path))
    if source_fmt == 'txt' and target_fmt == 'pdf':
        return _txt_to_pdf(source_path, str(out_path))
    if source_fmt == 'txt' and target_fmt == 'docx':
        return _txt_to_docx(source_path, str(out_path))
    if source_fmt in ('html', 'htm') and target_fmt == 'pdf':
        return _html_to_pdf(source_path, str(out_path))

    raise ConversionError(f'Конвертация {source_fmt} -> {target_fmt} не поддерживается')


def _pdf_to_docx(source: str, dest: str) -> str:
    if not PDF2DOCX_AVAILABLE:
        raise ConversionError('Установите pdf2docx: pip install pdf2docx')
    try:
        cv = Pdf2DocxConverter(source)
        cv.convert(dest)
        cv.close()
        return dest
    except Exception as e:
        raise ConversionError(f'Ошибка PDF->DOCX: {e}')


def _docx_to_pdf(source: str, dest: str) -> str:
    """DOCX -> PDF через reportlab (извлечение текста и создание PDF)."""
    if not DOCX_AVAILABLE or not REPORTLAB_AVAILABLE:
        raise ConversionError('Установите python-docx и reportlab')
    try:
        doc = Document(source)
        c = canvas.Canvas(dest, pagesize=A4)
        width, height = A4
        y = height - 50
        for para in doc.paragraphs:
            text = para.text
            if text.strip():
                c.setFont('Helvetica', 12)
                c.drawString(50, y, text[:100])
                y -= 18
            if y < 50:
                c.showPage()
                y = height - 50
        c.save()
        return dest
    except Exception as e:
        raise ConversionError(f'Ошибка DOCX->PDF: {e}')


def _txt_to_pdf(source: str, dest: str) -> str:
    if not REPORTLAB_AVAILABLE:
        raise ConversionError('Установите reportlab: pip install reportlab')
    try:
        with open(source, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        c = canvas.Canvas(dest, pagesize=A4)
        width, height = A4
        y = height - 50
        for line in text.splitlines():
            line = line[:120]
            if line.strip():
                c.setFont('Helvetica', 12)
                c.drawString(50, y, line)
                y -= 18
            if y < 50:
                c.showPage()
                y = height - 50
        c.save()
        return dest
    except Exception as e:
        raise ConversionError(f'Ошибка TXT->PDF: {e}')


def _txt_to_docx(source: str, dest: str) -> str:
    if not DOCX_AVAILABLE:
        raise ConversionError('Установите python-docx: pip install python-docx')
    try:
        with open(source, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        doc = Document()
        for line in text.splitlines():
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)
        doc.save(dest)
        return dest
    except Exception as e:
        raise ConversionError(f'Ошибка TXT->DOCX: {e}')


def _html_to_pdf(source: str, dest: str) -> str:
    if not XHTML2PDF_AVAILABLE:
        raise ConversionError('Установите xhtml2pdf: pip install xhtml2pdf')
    try:
        with open(source, 'r', encoding='utf-8', errors='ignore') as f:
            html = f.read()
        with open(dest, 'wb') as out:
            pisa_status = pisa.CreatePDF(html, dest=out, encoding='utf-8')
        if pisa_status.err:
            raise ConversionError('Ошибка создания PDF из HTML')
        return dest
    except ConversionError:
        raise
    except Exception as e:
        raise ConversionError(f'Ошибка HTML->PDF: {e}')
